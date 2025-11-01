from typing import List , Optional
import uuid 
from datetime import datetime
import PyPDF2
import io
import docx
from app.application.interfaces.document_repository import IDocumentRepositroy
from app.application.interfaces.embedding_service import IEmbeddingService
from app.application.interfaces.vector_store import IvectorStore
from app.application.interfaces.storage_service import IStorageService
from app.domain.entities.document import Document
from app.domain.exceptions import InvalidDocumentFormatError , DocumentNotFoundError

class DocumentService:
    def __init__(self  ,document_repo : IDocumentRepositroy,
                 embedding_service:IEmbeddingService,
                 vector_store : IvectorStore,
                 storage_service:IStorageService):
        self.document_repo = document_repo
        self.embedding_service = embedding_service
        self.vector_store = vector_store
        self.storage_service = storage_service

    async def process_document(
            self,
            filename:str,
            content:bytes,
            user_id : str
    ) -> Document:
        """
        Complete document processing pipeline:
        1. Extract text
        2. Create document entity
        3. Chunk text
        4. Generate embeddings
        5. Store in vector database
        6. Store original file
        7. Save metadata to database
        """
        
        # 1. Extract text
        text_content = await self._extract_text(content , filename)

        if not text_content.strip():
            raise InvalidDocumentFormatError("Document is empty or could not be read")
        
        # 2. Create document entity
        document = Document(
            id = str(uuid.uuid4()),
            filename=filename,
            content = text_content,
            created_at=datetime.utcnow(),
            user_id=user_id
        )

        # 3. split into chunks (business logic in entity) - Need more optimization later
        chunks = document.split_into_chunks(chunk_size=1000)

        # 4. create embeddings for all chunks
        embeddings = await self.embedding_service.create_embeddings_batch(chunks)

        # 5. Store chunks with embeddings in vector database
        for i , (chunk , embedding) in enumerate(zip(chunks , embeddings)):
            await self.vector_store.upsert(
                id=f"{document.id}_chunk_{i}",
                embedding=embedding,
                metadata={
                    "document_id":document.id,
                    "chunk_index" : i ,
                    "text" : chunk,
                    "user_id" :user_id,
                    "filename":filename

                }

            )

        # 6. Store original file in object storage
        storage_key = f"documents/{user_id}/{document.id}/{filename}"
        await self.storage_service.upload(
            key = storage_key,
            content=content
        )

        # 7. Save document metadata to database
        await self.document_repo.save(document)

        return document
    

    async def get_document(self,document_id :str , user_id:str) -> Document:
        """Get a document by ID"""
        document = await self.document_repo.get_by_id(document_id)

        if not document :
            raise DocumentNotFoundError(f"Document {document_id} not found")
        
        if document.user_id != user_id:
            raise PermissionError("Not authorized to access this document")
        
        return document
    

    async def list_documents(self , user_id:str , limit:int = 50 , offset:int = 0) -> List[Document]:
        """List user's documents"""
        return await self.document_repo.list_by_user(
            user_id=user_id,
            limit=limit ,
            offset=offset
        )
    
    async def search_documents(
            self, user_id:str , query:str , limit:int = 10
    ) -> List[Document]:
        """Search user's documents"""
        return await self.document_repo.search_by_user(
            user_id=user_id,
            query=query,
            limit=limit
        )
    
    async def delete_document(self , document_id:str , user_id:str) -> None:
        """Delete document completely:
        1. Delete from vector store
        2. Delete from object storage
        3. Delete from database
        """

        document = await self.get_document(document_id , user_id)

        # 1. Delete from vector store
        try:
            # Delete vector embeddings
            pass
        except Exception as e:
            print(f"Waring: Could not delete vectors for {document_id}: {e}")

        # 2. Delete from object storage
        storage_key = f"documents/{user_id}/{document_id}/{document.filename}"
        try:
            await self.storage_service.delete(key=storage_key)
        except Exception as e :
            print("Warning : could not delete file from storage")

        # 3. Delete from database
        await self.document_repo.delete(document_id)

    async def get_document_count(self , user_id :str) -> int:
        """Get total document count for user"""
        return await self.document_repo.count_by_user(user_id)
    

    async def _extract_text(self , content:bytes , filename:str) ->str:
        
        """Extract text from different file format"""
        file_extention = filename.split('.')[-1].lower()

        if file_extention == "pdf":
            return await self._extract_pdf_text(content)
        elif file_extention == "docx":
            return await self._extract_docx_text(content)
        elif file_extention in ["txt" , "md"] :
            return content.decode('utf-8')
        else:
            raise InvalidDocumentFormatError(
                f"Unsupported file format:{file_extention}"
            )

        


    async def _extract_pdf_text(self , content:bytes) ->str:
        """Extract text from pdf"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return text
        except Exception as e :
            raise InvalidDocumentFormatError(f"Could not read PDF: {str(e)}")

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text
        except Exception as e:
            raise InvalidDocumentFormatError(f"Could not read DOCX: {str(e)}")



        

        
        


        


        



